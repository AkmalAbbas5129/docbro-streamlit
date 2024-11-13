from langchain_openai import AzureChatOpenAI
import mammoth
from io import BytesIO
from markdownify import markdownify as md
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
import re
import os

os.environ["AZURE_OPENAI_API_KEY"] = st.secrets["openai_api_key"]
os.environ["AZURE_OPENAI_ENDPOINT"] = st.secrets["azure_endpoint"]
os.environ["AZURE_OPENAI_API_VERSION"] = st.secrets["api_version"]
os.environ["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"] = st.secrets["deployment_name"]
os.environ["OPENAI_MODEL_NAME"] = "gpt-4o"

def add_code_style(doc):
    """
    Adds a 'Code' style to the document with a monospaced font.
    """
    styles = doc.styles
    if 'Code' not in styles:
        code_style = styles.add_style('Code', 1)  # 1 is for paragraph style
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(10)  # Set font size for readability
        code_style.font.highlight_color = None


def add_content_to_doc(doc, content):
    """
    Adds content to the Word document based on Markdown-like headers and lists.
    """
    lines = content.splitlines()
    in_code_block = False  # Track if we're inside a code block

    for line in lines:
        line = line.strip()

        if not line:
            continue  # Skip empty lines

        # Handle code blocks (start and end with ```)
        if line.startswith("```"):
            in_code_block = not in_code_block
            if in_code_block:
                # Start of a code block
                doc.add_paragraph("Code Block:", style='Intense Quote')
            continue  # Skip the ``` line itself

        # Add code block content
        if in_code_block:
            code_paragraph = doc.add_paragraph(line)
            code_paragraph.style = 'Code'  # Apply a monospace font style if available
            continue

        # Check for headers based on the number of '#' characters
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            header_level = len(header_match.group(1))
            header_text = header_match.group(2)
            doc.add_heading(header_text, level=min(header_level, 4))  # Limit heading levels to 4
            continue

        # Check for bullet points or list items
        if re.match(r'^[-*]\s+(.*)', line):
            bullet_text = re.sub(r'^[-*]\s+', '', line)
            doc.add_paragraph(bullet_text, style='List Bullet')
            continue

        # Check if it’s a JSON object block (just for display, not parsing JSON)
        if line.startswith("{") or line.startswith("}"):
            json_paragraph = doc.add_paragraph(line)
            json_paragraph.style = 'Code'  # Apply a monospace font style if available
            continue

        # Add regular text as a paragraph
        doc.add_paragraph(line)


def create_word_doc_from_markdown(content: str, filename="MarkdownToWord.docx"):
    """
    Creates a Word document from markdown-like structured content.
    """
    doc = Document()

    # Add the custom Code style to the document
    add_code_style(doc)

    # Add title if provided as the first line with "###" markdown header
    first_line = content.splitlines()[0]
    if first_line.startswith("###"):
        title_text = first_line.strip('# ').strip()
        title = doc.add_heading(title_text, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        content = content.replace(first_line, '', 1)  # Remove title from the main content

    # Process and add the rest of the content
    add_content_to_doc(doc, content)

    # # Save the document
    # doc.save(filename)
    # Save the document to a BytesIO object
    word_file = BytesIO()
    doc.save(word_file)
    word_file.seek(0)
    return word_file
    # print(f"Document '{filename}' created successfully.")


def load_template(path):
    with open(path, 'r', encoding='utf-8') as file:
        brd_template_content = file.read()
    return brd_template_content.replace('\ufeff', '')

def get_llm():
    default_llm = AzureChatOpenAI(

        openai_api_version=st.secrets["api_version"],
        azure_deployment=st.secrets["deployment_name"],
        azure_endpoint=st.secrets["azure_endpoint"],
        api_key=st.secrets["openai_api_key"],
        model_name="gpt-4o",
        temperature=0,
        top_p=0.3
    )
    return default_llm


def read_docx_as_markdown_from_dir(file_path):
    # Open the .docx file from the local directory
    with open(file_path, "rb") as docx_file:
        docx_data = docx_file.read()

    # Convert the .docx bytes to HTML using Mammoth
    with BytesIO(docx_data) as docx_stream:
        result = mammoth.convert_to_html(docx_stream)
        html_output = result.value  # The converted HTML string

    # Convert HTML to Markdown using markdownify
    markdown_output = md(html_output)

    return markdown_output