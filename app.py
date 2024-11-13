from streamlit_agraph import agraph, Node, Edge, Config
from streamlit_option_menu import option_menu
from main_testcase import testcase_main
from main_Summary import summary_main
from streamlit_markmap import markmap
from main_testcase import testcase_main
from dotenv import load_dotenv
from main_BRD import brd_main
from main_SDD import sdd_main
from main_TSD import tsd_main
from utils import create_word_doc_from_markdown
from docx import Document
import streamlit as st
from io import BytesIO
import zipfile
import json
import os
from io import BytesIO
from langsmith import traceable
from pathlib import Path




langchain_api_key = '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'

os.environ["LANGCHAIN_API_KEY"] = langchain_api_key #os.getenv("LANGCHAIN_API_KEY")
os.environ["LANGCHAIN_PROJECT"] = 'docbro-prod'
os.environ["LANGCHAIN_TRACING_V2"]="true"
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"

# Function to list all files in the given directory
def list_files_in_directory(directory):
    files = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]
    return files


def list_folders_in_directory(directory_path):
    folders = [name for name in os.listdir(directory_path) if os.path.isdir(os.path.join(directory_path, name))]
    return folders


def list_files_in_directory_sub_fold(directory, extensions=None):
    if extensions is None:
        extensions = ['.py', '.js', '.html', '.css', '.java', '.cpp', '.ts']  # Add more extensions as needed

    coding_files = []
    for dirpath, dirnames, filenames in os.walk(directory):
        for filename in filenames:
            if not filename.startswith('.') and any(filename.endswith(ext) for ext in extensions):
                coding_files.append(os.path.join(dirpath, filename))

    return coding_files


def save_uploaded_zip(uploaded_file, save_directory):
    if not os.path.exists(save_directory):
        os.makedirs(save_directory)

    # Save the uploaded file to a buffer
    with BytesIO(uploaded_file.read()) as buffer:
        with zipfile.ZipFile(buffer, 'r') as zip_ref:
            zip_ref.extractall(save_directory)


@st.dialog("File Summary")
def file_summary(path):
    if 'summary_generated' not in st.session_state:
        st.session_state.summary_generated = False
        st.session_state.summary_result = ""

    if not st.session_state.summary_generated:
        with st.spinner(f"Generating Summary for {Path(path).stem}"):
            res = summary_main(path)
            st.session_state.summary_result = res
            st.session_state.summary_generated = True

    # Display the generated summary from session state
    st.markdown(f"<div class='stResult'>{st.session_state.summary_result}</div>", unsafe_allow_html=True)

    # Offer the summary result for download
    save_and_download_result_docs(f"{Path(path).stem}_Summary", st.session_state.summary_result)
#

# Custom CSS for styling
def add_custom_css():
    st.markdown(
        """
        <style>
        .main {
            background-color: #f0f2f6;
            padding: 2rem;
        }
        .stSidebar {
            background-color: #262B3F;
            padding: 1rem;
            color: #ffffff;
        }
        .stSidebarTitle {
            font-size: 5rem;
            font-weight: bold;
            color: #F26522;
            margin-bottom: 0rem;
        }
        .stTitle {
            font-size: 1.5rem;
            font-weight: bold;
            color: #262B3F;
            margin-bottom: 2rem;
        }
        [data-testid="stSidebar"] .stRadio > label {
        color: #FFFFFF; 
        }
        [data-testid="stSidebar"] .stRadio div[role='radiogroup'] label p {
            color: #F26522; 
        }
        [data-testid="stSidebar"] hr {
            border-color: #FFFFFF;
            border-width: 2px;  /* Optional: Change the thickness of the divider */
        }
        .custom-label {
            color: #FFFFFF;
            font-size: 16px;
            font-weight: bold;
        }
        .custom-label1 {
            color: #FFFFFF;
            font-size: 14px;
        }
        .stButton > button {
            background-color: #262B3F;
            color: #FFFFFF;
            padding: 10px 24px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            margin-top: 10px;
            text-align: center;
        }
        .stButton > button:hover {
            background-color: #FF6347;  /* Background color on hover */
            color: white;  /* Text color on hover */
        }
        .stHeader {
            color: #F26522;
            font-family: 'Arial', sans-serif;
        }
        .stDownloadButton > button {
            background-color: #262B3F;  /* Background color */
            color: white;  /* Text color */
            border-radius: 10px;  /* Optional: Rounded corners */
            border: none;  /* Optional: Remove border */
        }
        .stDownloadButton > button:hover {
            background-color: #FF6347;  /* Background color on hover */
            color: white;  /* Text color on hover */
        }
        .stResult {
            background-color: #e6ffee;
            border-left: 5px solid #4CAF50;
            padding: 10px;
            border-radius: 4px;
            margin-top: 10px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# Function to save uploaded file to a specific directory
def save_uploaded_file(directory, uploaded_file):
    if uploaded_file is not None:
        file_path = os.path.join(directory, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        return file_path
    return None


# Function to save and download the result as a Word document
def save_and_download_result(file_name, content):
    doc = Document()
    doc.add_paragraph(content)

    # Save to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download File",
        data=buffer,
        file_name=f"{file_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def save_and_download_result_docs(file_name, content):
    try:
        if content:
                word_file = create_word_doc_from_markdown(content)
                st.download_button(
                    label="Download File",
                    data=word_file,
                    file_name=f"{file_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Please enter text to convert.")
    except Exception as e:
        st.warning("Unable to convert Markdown to word Document.")

def generate_markdown(folder_path, level=1):
    """Recursively generate the Markdown structure for folder contents."""
    contents = os.listdir(folder_path)
    f_name = os.path.basename(folder_path)

    # Create a Markdown structure for the folder contents
    markdown_structure = f"{'#' * level} {f_name}\n"

    for item in contents:
        item_path = os.path.join(folder_path, item)

        if os.path.isdir(item_path):
            # For directories, recursively get the contents
            markdown_structure += generate_markdown(item_path, level + 1)
        else:
            # For files, add them to the markdown at the current level
            markdown_structure += f"{'#' * (level + 1)} {item}\n"

    return markdown_structure


def show_folder_contents(folder_path):
    st.markdown("""
                <h3 style='font-size:16px;'>Repository Structure</h3>
                """, unsafe_allow_html=True)

    markdown_structure = generate_markdown(folder_path)
    # st.write(markdown_structure)
    # print(markdown_structure)

    # Visualize the Markdown structure using markmap
    markmap(markdown_structure, height=400)


 
def add_table_and_column_nodes(table, table_color="orange", column_color="lightblue"):
    if table['name']:
        table_node = Node(
            id=table['name'],
            label=f"Table: {table['name']}",
            size=20,
            color=table_color,
            title=f"Table: {table['name']} - {len(table['columns'])} columns"
        )
        st.session_state["nodes"].append(table_node)
 
        for column in table['columns']:
            column_node = Node(
                id=f"{table['name']}_{column['name']}",
                label=f"Column: {column['name']} ({column['type']})",
                size=15,
                color=column_color,
                title=f"Column {column['name']} of type {column['type']}"
            )
            st.session_state["nodes"].append(column_node)
            st.session_state["edges"].append(Edge(source=table['name'], target=f"{table['name']}_{column['name']}", label="has_column"))
 
# Helper function to add nodes and edges for backend classes and methods
def add_class_and_method_nodes(klass, class_color="green", method_color="lightgreen"):
    if klass['name']:
        class_node = Node(
            id=klass['name'],
            label=f"Class: {klass['name']}",
            size=20,
            color=class_color,
            title=klass.get('description', None)
        )
        st.session_state["nodes"].append(class_node)
 
        for method in klass['methods']:
            method_node = Node(
                id=f"{klass['name']}_{method['name']}",
                label=f"Method: {method['name']}",
                size=15,
                color=method_color,
                title=method.get('description', None)
            )
            st.session_state["nodes"].append(method_node)
            st.session_state["edges"].append(Edge(source=klass['name'], target=f"{klass['name']}_{method['name']}", label="has_method"))
 
# Helper function to add nodes and edges for frontend components and API calls
def add_frontend_component_nodes(component, component_color="purple"):
    if component['name']:
        component_node = Node(
            id=component['name'],
            label=f"Frontend Component: {component['name']}",
            size=20,
            color=component_color,
            title=f"State Management: {component['stateManagement']}"
        )
        st.session_state["nodes"].append(component_node)
 
        for api_call in component['apiCalls']:
            api_node = Node(
                id=f"{component['name']}_{api_call['method']}",
                label=f"API: {api_call['method']} {api_call['endpoint']}",
                size=15,
                color="lightpurple",
                title=api_call.get('description', None)
            )
            st.session_state["nodes"].append(api_node)
            st.session_state["edges"].append(Edge(source=component['name'], target=f"{component['name']}_{api_call['method']}", label="makes_api_call"))
 
def visualize_project_structure(project_structure):
    # Clear nodes and edges in session state for fresh visualization each time
    st.session_state["nodes"].clear()
    st.session_state["edges"].clear()
 
    # Add Database tables and columns as nodes
    for table in project_structure['database']['tables']:
        try:
            add_table_and_column_nodes(table)
        except Exception as e:
            pass
 
    # Add backend classes and methods as nodes
    for klass in project_structure['backend']['classes']:
        try:
            add_class_and_method_nodes(klass)
        except Exception as e:
            pass
 
    # Add frontend components and API calls as nodes
    for component in project_structure['frontend']['components']:
        try:
            add_frontend_component_nodes(component)
        except Exception as e:
            pass
 
    # Add relationships for tables (many-to-one, one-to-many)
    for table in project_structure['database']['tables']:
        if "relations" in table and table["relations"]:
            for relation, related_tables in table["relations"].items():
                if isinstance(related_tables, list):
                    for related_tab in related_tables:
                        if isinstance(related_tab, dict):
                            for key, related_table in related_tab.items():
                                st.session_state["edges"].append(
                                    Edge(source=table['name'], target=related_table, label=relation))
                        else:
                            st.session_state["edges"].append(
                                Edge(source=table['name'], target=related_tab, label=relation))


    # Add relations for backend classes (uses/connectsTo)
    for klass in project_structure['backend']['classes']:
        if "relations" in klass and klass["relations"]:
            for relation, related_classes in klass["relations"].items():
                for related_class in related_classes:
                    st.session_state["edges"].append(Edge(source=klass['name'], target=related_class, label=relation))

    for comp in project_structure['frontend']['components']:
        if comp['parent']:
            st.session_state["edges"].append(Edge(source=comp['name'], target=comp['parent'], label="conented to"))


    # Configure the graph
    config = Config(
        height=500,
        width=700,
        nodeHighlightBehavior=True,
        highlightColor="#F7B7A6",
        directed=True,
        physics=True,
        collapsible=True,
        hierarchical=False
    )
 
    # Render the graph with nodes and edges from session state
    return_value = agraph(nodes=st.session_state["nodes"], edges=st.session_state["edges"], config=config)    
 
    # Optional: return the graph structure for further analysis or debugging
    return return_value


def main():
    add_custom_css()

    pages = ["BRD", "SDD", "TSD", "TestCases"]
    st.markdown(
        "<div class='stTitle'>Your Intelligent Assistant for Effortless Documentation.</div>",
        unsafe_allow_html=True)
    task = option_menu(None, pages,
                       icons=['file-earmark-text', "file-earmark-text", 'file-earmark-text', 'code'],
                       menu_icon="cast", default_index=0, orientation="horizontal")

    # Sidebar logo
    logo_path = "https://systemsltd.com/themes/custom/sysltd/SystemsLogo-02.svg"
    logo_path1 = "https://systemsltd.com/themes/custom/sysltd/SystemsLogo-01.svg"
    # Example of displaying a logo with a background color
    st.logo(logo_path, icon_image=logo_path1)

    # Title of the project
    st.sidebar.markdown(
        "<div class='stSidebarTitle'>DocBro</div>",
        unsafe_allow_html=True)

    st.sidebar.divider()


    if task == "BRD":

        # Hardcoded directories
        brd_template_dir = "./BRD-Docs/brd-template"
        rfp_docs_dir = "./BRD-Docs/rfp-document"
        transcriptions_dir = "./BRD-Docs/meeting-transcription"
        relevant_docs_dir = "./BRD-Docs/other"

        # File uploaders
        st.sidebar.markdown("<div class='custom-label'>Upload BRD Template:</div>", unsafe_allow_html=True)
        uploaded_brd_template = st.sidebar.file_uploader("Upload BRD Template:", type=["md", "docx"],
                                                         key="brd_template", label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload RFP Document:</div>", unsafe_allow_html=True)
        uploaded_rfp_doc = st.sidebar.file_uploader("Upload RFP Document:", type=["pdf"], key="rfp_document",
                                                    label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload Meeting Transcription:</div>", unsafe_allow_html=True)
        uploaded_transcription = st.sidebar.file_uploader("Upload Meeting Transcription:", type=["md", "docx"],
                                                          key="transcription", label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload Other Relevant Document:</div>", unsafe_allow_html=True)
        uploaded_relevant_docs = st.sidebar.file_uploader("Upload Other Relevant Document:", type=["md", "docx"],
                                                          key="relevant_docs", label_visibility="hidden")

        # Save uploaded files
        if uploaded_brd_template:
            brd_template_path = save_uploaded_file(brd_template_dir, uploaded_brd_template)
        if uploaded_rfp_doc:
            rfp_doc_path = save_uploaded_file(rfp_docs_dir, uploaded_rfp_doc)
        if uploaded_transcription:
            transcription_path = save_uploaded_file(transcriptions_dir, uploaded_transcription)
        if uploaded_relevant_docs:
            relevant_docs_path = save_uploaded_file(relevant_docs_dir, uploaded_relevant_docs)

        # List files in each directory
        brd_files = list_files_in_directory(brd_template_dir)
        rfp_files = list_files_in_directory(rfp_docs_dir)
        transcription_files = list_files_in_directory(transcriptions_dir)
        relevant_docs = list_files_in_directory(relevant_docs_dir)

        # Display files for selection
        st.markdown("""
                    <h3 style='font-size:16px;;'>Select the 📑files for BRD generation</h3>
                    <hr style='border:1px solid red; margin-top: 0px; margin-bottom: 5px;'>
                """, unsafe_allow_html=True)

        selected_brd_file = st.selectbox("BRD Template file:", brd_files, key="brd_file",
                                         label_visibility="hidden",
                                         index=None,
                                         placeholder="Select BRD Template",
                                         )

        selected_rfp_file = st.selectbox("RFP Document file:", rfp_files, key="brd_rfp_file",
                                         label_visibility="hidden",
                                         index=None,
                                         placeholder="Select RFP File",
                                         )
        selected_transcription_file = st.selectbox("Transcription file:", transcription_files,
                                                   key="brd_transcription_file",
                                                   label_visibility="hidden",
                                                   index=None,
                                                   placeholder="Select Meeting Transcription",
                                                   )
        selected_relevant_docs = st.selectbox("Relevant file:", relevant_docs,
                                              key="relevant_doc_file_brd",
                                              label_visibility="hidden",
                                              index=None,
                                              placeholder="Select Other Relevant Document",
                                              )

        # Submit button
        if st.button("Generate BRD"):
            if selected_brd_file:
                if selected_rfp_file or selected_transcription_file or selected_relevant_docs:
                    with st.spinner('Generating BRD...'):
                        brd_file_path = os.path.join(brd_template_dir, selected_brd_file) if selected_brd_file else ""
                        rfp_file_path = os.path.join(rfp_docs_dir, selected_rfp_file) if selected_rfp_file else ""
                        transcription_file_path = os.path.join(transcriptions_dir,
                                                               selected_transcription_file) if selected_transcription_file else ""
                        relevant_file_path = os.path.join(relevant_docs_dir,
                                                          selected_relevant_docs) if selected_relevant_docs else ""

                        # brd_file_path = brd_template_dir + '/' + selected_brd_file
                        # rfp_file_path = rfp_docs_dir + '/' + selected_rfp_file
                        # transcription_file_path = transcriptions_dir + '/' + selected_transcription_file
                        # if not selected_relevant_docs:
                        #     relevant_file_path = ""
                        # else:
                        #     relevant_file_path = relevant_docs_dir + '/' + selected_relevant_docs

                        # Process the selected files
                        brd_result = brd_main(brd_file_path, rfp_file_path, transcription_file_path, relevant_file_path)

                        # brd_result = "Generated BRD"
                        st.markdown(f"<div class='stResult'>{brd_result}</div>", unsafe_allow_html=True)
                        save_and_download_result("BRD_Result", brd_result)
                else:
                    st.error("Please upload one required files.")
            else:
                st.error("Please upload BRD Template.")

    elif task == "SDD":

        # Hardcoded directories
        sdd_template_dir = "./SDD-Docs/sdd-template"
        rfp_docs_dir = "./BRD-Docs/rfp-document"
        sdd_transcriptions_dir = "./SDD-Docs/sdd-meeting-transcription"
        generated_brd_dir = "./BRD-Docs/generated-brd"
        relevant_docs_dir = "./SDD-Docs/other"

        # File uploaders
        st.sidebar.markdown("<div class='custom-label'>Upload SDD Template:</div>", unsafe_allow_html=True)
        uploaded_sdd_template = st.sidebar.file_uploader("Upload SDD Template:", type=["md", "docx"],
                                                         key="sdd_template", label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload Generated BRD:</div>", unsafe_allow_html=True)
        uploaded_generated_brd = st.sidebar.file_uploader("Upload Generated BRD:", type=["md", "docx"],
                                                          key="generated_brd", label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload RFP Document:</div>", unsafe_allow_html=True)
        uploaded_rfp_doc = st.sidebar.file_uploader("Upload RFP Document:", type=["pdf"], key="rfp_document_sdd",
                                                    label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload Meeting Transcription:</div>", unsafe_allow_html=True)
        uploaded_transcription = st.sidebar.file_uploader("Upload Meeting Transcription:", type=["md", "docx"],
                                                          key="transcription_sdd", label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload Other Relevant Document:</div>", unsafe_allow_html=True)
        uploaded_relevant_docs = st.sidebar.file_uploader("Upload Other Relevant Document:", type=["md", "docx"],
                                                          key="relevant_docs_sdd", label_visibility="hidden")

        # Save uploaded files
        if uploaded_sdd_template:
            sdd_template_path = save_uploaded_file(sdd_template_dir, uploaded_sdd_template)
        if uploaded_generated_brd:
            generated_brd_path = save_uploaded_file(generated_brd_dir, uploaded_generated_brd)
        if uploaded_rfp_doc:
            rfp_doc_path = save_uploaded_file(rfp_docs_dir, uploaded_rfp_doc)
        if uploaded_transcription:
            transcription_path = save_uploaded_file(sdd_transcriptions_dir, uploaded_transcription)
        if uploaded_relevant_docs:
            relevant_docs_path = save_uploaded_file(relevant_docs_dir, uploaded_relevant_docs)

        # List files in each directory
        sdd_template_files = list_files_in_directory(sdd_template_dir)
        generated_brd_files = list_files_in_directory(generated_brd_dir)
        rfp_files = list_files_in_directory(rfp_docs_dir)
        transcription_files = list_files_in_directory(sdd_transcriptions_dir)
        relevant_docs = list_files_in_directory(relevant_docs_dir)

        # Display files for selection
        st.markdown("""
                    <h3 style='font-size:16px;;'>Select the 📑files for SDD generation</h3>
                    <hr style='border:1px solid red; margin-top: 0px; margin-bottom: 5px;'>
                """, unsafe_allow_html=True)
        selected_sdd_template_file = st.selectbox("SDD Template file:", sdd_template_files, key="sdd_template_files",
                                                  label_visibility="hidden",
                                                  index=None,
                                                  placeholder="Select SDD Template",
                                                  )
        selected_brd_file = st.selectbox("Generated BRD file:", generated_brd_files, key="generated_brd_files",
                                         label_visibility="hidden",
                                         index=None,
                                         placeholder="Select BRD",
                                         )
        selected_rfp_file = st.selectbox("RFP Document file:", rfp_files, key="sdd_rfp_file",
                                         label_visibility="hidden",
                                         index=None,
                                         placeholder="Select RFP Document",
                                         )
        selected_transcription_file = st.selectbox("Transcription file:", transcription_files,
                                                   key="sdd_transcription_file",
                                                   label_visibility="hidden",
                                                   index=None,
                                                   placeholder="Select SDD Meeting Transcription",
                                                   )
        selected_relevant_docs = st.selectbox("Relevant file:", relevant_docs,
                                              key="relevant_doc_file_sdd",
                                              label_visibility="hidden",
                                              index=None,
                                              placeholder="Select Other Relevant Document",
                                              )

        # Submit button
        if st.button("Generate SDD"):
            if selected_sdd_template_file:
                if selected_transcription_file or selected_rfp_file or selected_brd_file or selected_relevant_docs:
                    with st.spinner('Generating SDD...'):
                        sdd_template_file_path = os.path.join(sdd_template_dir,
                                                              selected_sdd_template_file) if selected_sdd_template_file else ""
                        rfp_file_path = os.path.join(rfp_docs_dir, selected_rfp_file) if selected_rfp_file else ""
                        transcription_file_path = os.path.join(sdd_transcriptions_dir,
                                                               selected_transcription_file) if selected_transcription_file else ""
                        generated_brd_file_path = os.path.join(generated_brd_dir,
                                                               selected_brd_file) if selected_brd_file else ""
                        relevant_file_path = os.path.join(relevant_docs_dir,
                                                          selected_relevant_docs) if selected_relevant_docs else ""

                        # sdd_template_file_path = sdd_template_dir + '/' + selected_sdd_template_file
                        # rfp_file_path = rfp_docs_dir + '/' + selected_rfp_file
                        # transcription_file_path = sdd_transcriptions_dir + '/' + selected_transcription_file
                        # generated_brd_file_path = generated_brd_dir + '/' + selected_brd_file
                        # if not selected_relevant_docs:
                        #     relevant_file_path = ""
                        # else:
                        #     relevant_file_path = relevant_docs_dir + '/' + selected_relevant_docs

                        # Process the selected files
                        sdd_result = sdd_main(sdd_template_file_path, generated_brd_file_path, rfp_file_path,
                                              transcription_file_path, relevant_file_path)

                        # sdd_result = "SDD Generated"
                        st.markdown(f"<div class='stResult'>{sdd_result}</div>", unsafe_allow_html=True)
                        save_and_download_result("SDD_Result", sdd_result)
                else:
                    st.error("Please upload one required files.")
            else:
                st.error("Please upload SDD Template.")


    elif task == "TSD":

        # # Hardcoded directories
        repositories_dir = "./repos"

        st.sidebar.markdown("<div class='custom-label'>Upload a ZIP file containing your folder:</div>",
                            unsafe_allow_html=True)
        uploaded_repo = st.sidebar.file_uploader("Upload a ZIP file containing your folder", type=["zip"],
                                                 label_visibility="hidden")

        if uploaded_repo:
            save_directory = "./repos"
            save_uploaded_zip(uploaded_repo, save_directory)
            st.sidebar.success(f"ZIP file unzipped and saved in {save_directory}")

        repositories = list_folders_in_directory(repositories_dir)

        # Dropdown for selecting a repository
        st.markdown("""
            <h3 style='font-size:16px;'>Select the Project Repository</h3>
            <hr style='border:1px solid red; margin-top: 0px; margin-bottom: 5px;'>
        """, unsafe_allow_html=True)

        selected_repo = st.selectbox("Select Project Repository (codebase):", repositories,
                                     key="tsd_repo",
                                     label_visibility="hidden",
                                     index=None,
                                     placeholder="Select Project Repository (codebase)")

        if selected_repo:
            show_folder_contents(repositories_dir + '/' + selected_repo)
            tab1, tab2 = st.tabs(["Detailed TSD", "File Summary"])

            with tab1:
                if "nodes" not in st.session_state:
                    st.session_state["nodes"] = []
                if "edges" not in st.session_state:
                    st.session_state["edges"] = []


                final_summary = ""
                with st.form("visualization_form"):
                    st.title(f"{selected_repo.upper()} Detailed Explaination")

                    submit_visualization = st.form_submit_button(f"Explain {selected_repo} Code")
                    if submit_visualization:

                        # # user_input = st.text_area("Prompt:",prompt_template , height=200)
                        # if st.button(f"{selected_repo} Explain Code"):
                        #     # st.write(user_input)
                            
                        with st.spinner('Generating TSD...'):
                            tsd_repo_path = os.path.join(repositories_dir, selected_repo) if selected_repo else ""
                            # Process the selected files
                            prompt = ""

                            final_summary = tsd_main(tsd_repo_path, prompt)

                            st.markdown(f"<div class='stResult'>{final_summary}</div>", unsafe_allow_html=True)

                            

                            try:
                                tsd_json = json.loads(final_summary.split("```")[1].strip("json "))
                                
                                st.markdown("<hr style='border:2px solid orange; margin:10px 0;'>",
                                            unsafe_allow_html=True)  # Green line with 2px thickness
                                st.write("### Node Types:")
                                st.write("🟩 **Backend Classes**: Backend Classes of the project.")
                                st.write("🟧 **Database Tables**: Represents database tables.")
                                st.write("🟪 **Frontend Components**: Frontend components in the project.")

                                visualize_project_structure(tsd_json)
                            except Exception as e:
                                st.markdown(f"<div class='stResult'>Invalid Json, Unable to visualize.</div>",
                                            unsafe_allow_html=True)
                                
                
                if submit_visualization:                
                    save_and_download_result_docs(f"{selected_repo}_TSD", final_summary)


            with tab2:
                # Main App Logic
                st.header("File Summary")

                # Session state to control dialog open/close and selected file
                if 'dialog_open' not in st.session_state:
                    st.session_state.dialog_open = False

                # Simulating listing of repository files
                repo_files = list_files_in_directory_sub_fold(os.path.join(repositories_dir, selected_repo))

                # Display files in the selected repository
                st.markdown("""
                    <h3 style='font-size:16px;'>Files in Selected Repository</h3>
                    <hr style='border:1px solid blue; margin-top: 0px; margin-bottom: 5px;'>
                """, unsafe_allow_html=True)

                # List repository files and provide "Generate Summary" buttons
                for file in repo_files:
                    st.markdown(f"<div style='font-size:14px;'>{file}</div>", unsafe_allow_html=True)

                    # Generate Summary Button
                    if st.button(f"Generate Summary", key=f"summary_{file}"):
                        st.session_state.dialog_open = True
                        st.session_state.summary_generated = False  # Reset to ensure generation
                        file_summary(file)

                    st.markdown("<hr>", unsafe_allow_html=True)


    elif task == "TestCases":

        # Hardcoded directories
        generated_brd_dir = "./BRD-Docs/generated-brd"
        generated_sdd_dir = "./SDD-Docs/generated-sdd"
        repositories_dir = "./repos"

        # File uploaders
        st.sidebar.markdown("<div class='custom-label'>Upload BRD Document:</div>", unsafe_allow_html=True)
        uploaded_generated_brd = st.sidebar.file_uploader("Upload BRD Document:", type=["md", "docx"], key="tc_brd",
                                                          label_visibility="hidden")
        st.sidebar.markdown("<div class='custom-label'>Upload SDD Document:</div>", unsafe_allow_html=True)
        uploaded_generated_sdd = st.sidebar.file_uploader("Upload SDD Document:", type=["md", "docx"], key="tc_sdd",
                                                          label_visibility="hidden")

        st.sidebar.markdown("<div class='custom-label'>Upload a ZIP file containing your folder:</div>",
                            unsafe_allow_html=True)
        uploaded_repo = st.sidebar.file_uploader("Upload a ZIP file containing your folder", type=["zip"],
                                                 label_visibility="hidden")

        if uploaded_repo:
            save_directory = "./repos"
            save_uploaded_zip(uploaded_repo, save_directory)
            st.sidebar.success(f"ZIP file unzipped and saved in {save_directory}")

        # Save uploaded files
        if uploaded_generated_brd:
            generated_brd_path = save_uploaded_file(generated_brd_dir, uploaded_generated_brd)
        if uploaded_generated_sdd:
            generated_sdd_path = save_uploaded_file(generated_sdd_dir, uploaded_generated_sdd)

        # List files in each directory
        generated_brd_files = list_files_in_directory(generated_brd_dir)
        generated_sdd_files = list_files_in_directory(generated_sdd_dir)
        repositories = list_folders_in_directory(repositories_dir)

        # Display files for selection
        st.markdown("""
                    <h3 style='font-size:16px;;'>Select the 📑files for TestCase generation</h3>
                    <hr style='border:1px solid red; margin-top: 0px; margin-bottom: 5px;'>
                """, unsafe_allow_html=True)
        selected_brd_file = st.selectbox("Select BRD file:", generated_brd_files, key="generated_brd_files",
                                         label_visibility="hidden",
                                         index=None,
                                         placeholder="Select BRD file",
                                         )
        selected_sdd_file = st.selectbox("Select SDD file:", generated_sdd_files, key="generated_sdd_files",
                                         label_visibility="hidden",
                                         index=None,
                                         placeholder="Select SDD file",
                                         )

        selected_repo = st.selectbox("Select Project Repository (codebase):", repositories,
                                     key="tsd_repo",
                                     label_visibility="hidden",
                                     index=None,
                                     placeholder="Select Project Repository (codebase)",
                                     )

        # Submit button
        if st.button("Generate TestCases"):
            if selected_repo:
                with st.spinner('Generating TestCases...'):
                    generated_brd_file_path = os.path.join(generated_brd_dir,
                                                           selected_brd_file) if selected_brd_file else ""
                    generated_sdd_file_path = os.path.join(generated_sdd_dir,
                                                           selected_sdd_file) if selected_sdd_file else ""
                    tc_repo_path = os.path.join(repositories_dir, selected_repo) if selected_repo else ""

                   
                    # Process the selected files
                    tc_result = testcase_main(generated_brd_file_path, generated_sdd_file_path, tc_repo_path)

                    # tc_result= "Test cases Generated"

                    st.markdown(f"<div class='stResult'>{tc_result}</div>", unsafe_allow_html=True)
                    save_and_download_result("TSD_Result", tc_result)
            else:
                st.error("Please upload all required files.")


if __name__ == "__main__":
    main()
